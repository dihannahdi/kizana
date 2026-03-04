#!/usr/bin/env python3
"""
Process Bahtsul Masail documents (PDF, DOCX, DOC) and store in SQLite with FTS5.
Creates produk_hukum.sqlite with full-text search capabilities.
"""

import os
import sys
import sqlite3
import hashlib
import re
from pathlib import Path
from datetime import datetime

# Fix Windows console encoding for Arabic filenames
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

# ─── Configuration ───
BASE_DIR = Path(r"D:\nahdi\bahtsulmasail\Bahstul Masail Database")
OUTPUT_DB = Path(r"D:\nahdi\bahtsulmasail\backend\produk_hukum.sqlite")

def setup_database(db_path):
    """Create SQLite database with FTS5 for full-text search."""
    conn = sqlite3.connect(str(db_path))
    conn.execute("PRAGMA journal_mode = WAL")
    conn.execute("PRAGMA synchronous = NORMAL")
    
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT NOT NULL,
            subcategory TEXT DEFAULT '',
            source_file TEXT NOT NULL,
            file_type TEXT NOT NULL,
            file_size INTEGER DEFAULT 0,
            file_hash TEXT DEFAULT '',
            content TEXT NOT NULL,
            page_count INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        
        CREATE TABLE IF NOT EXISTS categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            doc_count INTEGER DEFAULT 0
        );
        
        -- FTS5 virtual table for full-text search
        CREATE VIRTUAL TABLE IF NOT EXISTS documents_fts USING fts5(
            title,
            content,
            category,
            content='documents',
            content_rowid='id',
            tokenize='unicode61'
        );
        
        -- Triggers to keep FTS in sync
        CREATE TRIGGER IF NOT EXISTS documents_ai AFTER INSERT ON documents BEGIN
            INSERT INTO documents_fts(rowid, title, content, category) 
            VALUES (new.id, new.title, new.content, new.category);
        END;
        
        CREATE TRIGGER IF NOT EXISTS documents_ad AFTER DELETE ON documents BEGIN
            INSERT INTO documents_fts(documents_fts, rowid, title, content, category) 
            VALUES('delete', old.id, old.title, old.content, old.category);
        END;
    """)
    
    conn.commit()
    return conn


def extract_pdf_text(filepath):
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(filepath))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages), len(reader.pages)
    except Exception as e:
        print(f"  [WARN] pypdf failed for {filepath.name}: {e}")
        # Fallback: try pdfplumber
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(str(filepath)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text.strip())
                return "\n\n".join(pages), len(pdf.pages)
        except Exception as e2:
            print(f"  [ERROR] pdfplumber also failed: {e2}")
            return "", 0


def extract_docx_text(filepath):
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = []
        for para in doc.paragraphs:
            try:
                text = para.text
                if text and text.strip():
                    paragraphs.append(text.strip())
            except Exception:
                continue
        # Also extract from tables
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        try:
                            text = cell.text
                            if text and text.strip():
                                paragraphs.append(text.strip())
                        except Exception:
                            continue
        except Exception:
            pass
        return "\n\n".join(paragraphs), len(paragraphs)
    except Exception as e:
        print(f"  [ERROR] docx failed for {filepath.name}: {e}")
        return "", 0


def extract_doc_text(filepath):
    """Extract text from DOC using antiword or textract."""
    try:
        import subprocess
        result = subprocess.run(
            ['antiword', str(filepath)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip(), result.stdout.count('\n')
    except Exception:
        pass
    
    # Fallback: try reading as binary and extracting text-like content
    try:
        with open(filepath, 'rb') as f:
            raw = f.read()
        # Simple heuristic: extract runs of printable chars
        text = raw.decode('utf-8', errors='ignore')
        # Filter to keep only readable portions
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 10]
        return "\n".join(lines[:500]), len(lines)
    except Exception as e:
        print(f"  [ERROR] doc extraction failed for {filepath.name}: {e}")
        return "", 0


def clean_title(filename):
    """Clean up filename to a readable title."""
    title = filename
    # Remove file extension
    title = re.sub(r'\.(pdf|docx|doc)$', '', title, flags=re.IGNORECASE)
    # Remove leading numbers and dots
    title = re.sub(r'^\d+[\.\)\s]+', '', title)
    # Clean up extra whitespace
    title = re.sub(r'\s+', ' ', title).strip()
    # Remove common prefixes
    title = re.sub(r'^(HASIL\s+)?KEPUTUSAN\s+', 'Keputusan ', title)
    return title if title else filename


def get_category(filepath, base_dir):
    """Determine category from directory structure."""
    rel = filepath.relative_to(base_dir)
    parts = rel.parts
    if len(parts) > 1:
        return parts[0]  # First directory level = category
    return "Umum"  # Root-level files


def get_subcategory(filepath, base_dir):
    """Get subcategory from deeper directory levels."""
    rel = filepath.relative_to(base_dir)
    parts = rel.parts
    if len(parts) > 2:
        return parts[1]  # Second directory level = subcategory
    return ""


def file_hash(filepath):
    """Compute MD5 hash of file."""
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        while chunk := f.read(8192):
            h.update(chunk)
    return h.hexdigest()


def process_all_documents(base_dir, conn):
    """Process all PDF, DOCX, and DOC files recursively."""
    extensions = {'.pdf', '.docx', '.doc'}
    files = []
    
    for root, dirs, filenames in os.walk(base_dir):
        for fname in filenames:
            ext = os.path.splitext(fname)[1].lower()
            if ext in extensions:
                files.append(Path(root) / fname)
    
    print(f"Found {len(files)} documents to process")
    
    # Track categories
    category_counts = {}
    success = 0
    failed = 0
    skipped = 0
    
    for i, filepath in enumerate(files):
        rel_path = str(filepath.relative_to(base_dir))
        print(f"[{i+1}/{len(files)}] Processing: {rel_path}")
        
        # Check if already processed (by file hash)
        fhash = file_hash(filepath)
        existing = conn.execute(
            "SELECT id FROM documents WHERE file_hash = ?", (fhash,)
        ).fetchone()
        if existing:
            print(f"  [SKIP] Already processed")
            skipped += 1
            continue
        
        # Extract text
        ext = filepath.suffix.lower()
        if ext == '.pdf':
            content, page_count = extract_pdf_text(filepath)
            file_type = 'pdf'
        elif ext == '.docx':
            content, page_count = extract_docx_text(filepath)
            file_type = 'docx'
        elif ext == '.doc':
            content, page_count = extract_doc_text(filepath)
            file_type = 'doc'
        else:
            continue
        
        if not content or len(content.strip()) < 20:
            print(f"  [WARN] No content extracted, skipping")
            failed += 1
            continue
        
        # Metadata
        title = clean_title(filepath.name)
        category = get_category(filepath, base_dir)
        subcategory = get_subcategory(filepath, base_dir)
        file_size = filepath.stat().st_size
        
        # Insert
        conn.execute("""
            INSERT INTO documents (title, category, subcategory, source_file, file_type, 
                                   file_size, file_hash, content, page_count)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (title, category, subcategory, rel_path, file_type, 
              file_size, fhash, content, page_count))
        
        # Track category
        category_counts[category] = category_counts.get(category, 0) + 1
        success += 1
        print(f"  [OK] {len(content)} chars, {page_count} pages")
    
    conn.commit()
    
    # Update categories table
    for cat, count in category_counts.items():
        conn.execute("""
            INSERT INTO categories (name, doc_count) VALUES (?, ?)
            ON CONFLICT(name) DO UPDATE SET doc_count = doc_count + excluded.doc_count
        """, (cat, count))
    conn.commit()
    
    print(f"\n{'='*60}")
    print(f"Processing complete!")
    print(f"  Success: {success}")
    print(f"  Failed:  {failed}")
    print(f"  Skipped: {skipped}")
    print(f"  Total:   {len(files)}")
    print(f"\nCategories:")
    for cat, count in sorted(category_counts.items()):
        print(f"  {cat}: {count} documents")


def main():
    print(f"Bahtsul Masail Document Processor")
    print(f"Base dir: {BASE_DIR}")
    print(f"Output DB: {OUTPUT_DB}")
    print(f"{'='*60}")
    
    if not BASE_DIR.exists():
        print(f"ERROR: Base directory not found: {BASE_DIR}")
        sys.exit(1)
    
    # Ensure output directory exists
    OUTPUT_DB.parent.mkdir(parents=True, exist_ok=True)
    
    # Setup database
    conn = setup_database(OUTPUT_DB)
    
    # Process documents
    process_all_documents(BASE_DIR, conn)
    
    # Print stats
    total = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
    total_chars = conn.execute("SELECT SUM(LENGTH(content)) FROM documents").fetchone()[0] or 0
    print(f"\nDatabase stats:")
    print(f"  Total documents: {total}")
    print(f"  Total text: {total_chars:,} characters ({total_chars/1_000_000:.1f} MB)")
    print(f"  DB size: {OUTPUT_DB.stat().st_size / 1_000_000:.1f} MB")
    
    # Test FTS search
    print(f"\nTesting FTS search for 'nikah'...")
    results = conn.execute("""
        SELECT d.id, d.title, d.category, snippet(documents_fts, 1, '<b>', '</b>', '...', 30)
        FROM documents_fts f
        JOIN documents d ON d.id = f.rowid
        WHERE documents_fts MATCH 'nikah'
        ORDER BY rank
        LIMIT 5
    """).fetchall()
    for r in results:
        print(f"  [{r[0]}] {r[1]} ({r[2]})")
        print(f"    {r[3][:100]}...")
    
    conn.close()
    print(f"\nDone! Database saved to: {OUTPUT_DB}")


if __name__ == "__main__":
    main()
