#!/usr/bin/env python3
"""
Kizana — Knowledge-Based Systems (KBS) Submission Package Generator
====================================================================
Generates all required files for Elsevier KBS journal submission:
  1. Cover Letter (Word)
  2. Title Page with Author Details (Word)
  3. Highlights (Word)
  4. Declaration of Competing Interests (Word) ← "Author Agreement"
  5. CRediT Author Statement (Word)
  6. Declaration of Generative AI Use (Word)

Usage:
    python generate_kbs_submission.py
"""

import os
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text, size=12, bold=False, italic=False,
         align=None, space_after=6, space_before=0, indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, size=14, bold=True, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def divider(doc):
    p = doc.add_paragraph("─" * 70)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(150, 150, 150)


# ============================================================
# AUTHOR & PAPER METADATA
# ============================================================
AUTHOR_NAME      = "Dihan Nahdi"
AUTHOR_EMAIL     = "nahdi@bahtsulmasail.tech"
AUTHOR_ORCID     = "[To be added — register at orcid.org]"
AFFILIATION      = "Independent Researcher, Kizana Search Project"
ADDRESS          = "Indonesia"
PAPER_TITLE      = ("Kizana: Cross-Lingual Information Retrieval for Classical "
                    "Islamic Jurisprudence Texts Using Domain-Specific Query Translation")
JOURNAL_NAME     = "Knowledge-Based Systems"
DATE             = "March 4, 2026"
GITHUB_URL       = "https://github.com/dihannahdi/kizana"


# ============================================================
# 1. COVER LETTER
# ============================================================
def create_cover_letter(out: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Header
    para(doc, AUTHOR_NAME,  bold=True, size=12, space_after=2)
    para(doc, f"{AFFILIATION} | {ADDRESS}", size=11, italic=True, space_after=2)
    para(doc, AUTHOR_EMAIL, size=11, space_after=2)
    para(doc, DATE, size=11, space_after=16)

    para(doc, "Editor-in-Chief", bold=True, size=12, space_after=2)
    para(doc, "Knowledge-Based Systems", italic=True, size=12, space_after=2)
    para(doc, "Elsevier", size=12, space_after=16)

    para(doc, "Dear Editor-in-Chief,", size=12, space_after=10)

    para(doc, (
        f"I am pleased to submit our manuscript entitled \"{PAPER_TITLE}\" "
        "for consideration for publication in Knowledge-Based Systems. "
        "This work presents a fully operational information retrieval system "
        "that applies knowledge-based techniques to enable cross-lingual access "
        "to the largest publicly-indexed corpus of classical Arabic Islamic "
        "scholarship: 7,872 books spanning 14 centuries of jurisprudence, "
        "theology, and legal discourse."
    ), size=12, space_after=8)

    para(doc, "This manuscript is directly aligned with KBS's core scope:", size=12, space_after=6, bold=True)

    contributions = [
        ("Knowledge representation & engineering: ",
         "We design and implement a rule-based domain-specific knowledge base "
         "of 80+ Islamic jurisprudence term mappings across Arabic, Indonesian, "
         "English, and transliterated forms — encoding centuries of terminological "
         "relationships as structured multilingual knowledge."),
        ("Intelligent retrieval system: ",
         "Kizana integrates BM25 lexical search with a domain-aware query "
         "translation layer and configurable scoring, operating entirely offline "
         "with zero API dependency — suitable for knowledge-based decision support "
         "in resource-constrained Islamic educational institutions (pesantren)."),
        ("Rigorous empirical evaluation: ",
         "We evaluate on a gold standard of 96 expert-annotated queries across "
         "4 languages (Indonesian, English, Arabic, mixed) and 4 legal domains, "
         "with inter-annotator agreement κ = 0.793. Kizana achieves MAP = 0.790, "
         "reaching 98.9% of the Google Translate + BM25 external baseline "
         "(MAP = 0.799, p = 0.148), with no statistically significant difference."),
        ("Ablation study: ",
         "A 10-configuration component ablation reveals that the query translation "
         "knowledge base contributes ΔMAP = +0.601 — the single largest improvement "
         "component — while morphological stemming and multi-variant expansion "
         "reduce precision, providing actionable insights for knowledge-based IR design."),
    ]

    for label, text in contributions:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.4)
        r1 = p.add_run(label)
        set_font(r1, bold=True, size=12)
        r2 = p.add_run(text)
        set_font(r2, size=12)

    para(doc, "", space_after=4)
    para(doc, (
        "The system serves the Indonesian Muslim community (250M+ speakers) "
        "and is live at bahtsulmasail.tech. The complete source code, evaluation "
        f"scripts, and gold standard dataset are publicly available at {GITHUB_URL}."
    ), size=12, space_after=8)

    para(doc, (
        "This manuscript has not been published previously, is not under "
        "consideration elsewhere, and all authors have approved its submission. "
        "I declare no competing interests."
    ), size=12, space_after=12)

    para(doc, "Thank you for considering this manuscript.", size=12, space_after=16)
    para(doc, "Sincerely,", size=12, space_after=8)
    para(doc, AUTHOR_NAME, bold=True, size=12, space_after=2)
    para(doc, AFFILIATION, italic=True, size=12, space_after=2)
    para(doc, AUTHOR_EMAIL, size=12, space_after=0)

    path = out / "1_cover_letter_kbs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


# ============================================================
# 2. TITLE PAGE
# ============================================================
def create_title_page(out: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    heading(doc, "TITLE PAGE", size=14, bold=True)
    divider(doc)

    heading(doc, "Title", size=12, bold=True, space_before=10)
    para(doc, PAPER_TITLE, size=12, italic=True, space_after=12)

    heading(doc, "Author", size=12, bold=True, space_before=8)
    para(doc, AUTHOR_NAME, size=12, space_after=2)
    para(doc, AFFILIATION, italic=True, size=12, space_after=2)
    para(doc, ADDRESS, size=12, space_after=2)
    para(doc, f"Email: {AUTHOR_EMAIL}", size=12, space_after=2)
    para(doc, f"ORCID: {AUTHOR_ORCID}", size=12, space_after=12)

    heading(doc, "Corresponding Author", size=12, bold=True, space_before=8)
    para(doc, AUTHOR_NAME, size=12, space_after=2)
    para(doc, f"Email: {AUTHOR_EMAIL}", size=12, space_after=12)

    divider(doc)
    heading(doc, "CRediT Author Statement", size=12, bold=True, space_before=10)
    para(doc, f"{AUTHOR_NAME}:", bold=True, size=12, space_after=4)

    credit_roles = [
        ("Conceptualization:", "Identified the research gap in cross-lingual access to classical Islamic texts."),
        ("Data curation:", "Compiled, annotated, and validated the 96-query gold standard dataset across 4 languages and 4 domains."),
        ("Formal analysis:", "Conducted all statistical analyses including paired t-tests, Cohen's d, inter-annotator agreement (κ, Krippendorff's α), MAP, MRR, and NDCG calculations."),
        ("Investigation:", "Designed and executed all 10-configuration ablation experiments and the Google Translate external baseline comparison."),
        ("Methodology:", "Developed the rule-based query translation layer, BM25 scoring enhancements, and the evaluation framework."),
        ("Project administration:", "Managed all aspects of system development, evaluation, and manuscript preparation."),
        ("Resources:", "Deployed and maintains the production system at bahtsulmasail.tech."),
        ("Software:", "Implemented the full Rust/Actix-web backend, SvelteKit frontend, query translation engine, and all evaluation scripts."),
        ("Validation:", "Validated gold standard annotations with inter-annotator agreement protocols."),
        ("Visualization:", "Produced all 6 publication-quality figures at 300 DPI."),
        ("Writing – original draft:", "Wrote the complete manuscript."),
        ("Writing – review and editing:", "Reviewed and revised all sections."),
    ]

    for role, desc in credit_roles:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.3)
        r1 = p.add_run(role + " ")
        set_font(r1, bold=True, size=11)
        r2 = p.add_run(desc)
        set_font(r2, size=11)

    divider(doc)
    heading(doc, "Funding Statement", size=12, bold=True, space_before=10)
    para(doc, (
        "This research did not receive any specific grant from funding agencies "
        "in the public, commercial, or not-for-profit sectors."
    ), size=12, space_after=12)

    divider(doc)
    heading(doc, "Acknowledgements", size=12, bold=True, space_before=10)
    para(doc, (
        "The author thanks the Indonesian pesantren community and Islamic scholars "
        "whose accumulated knowledge forms the intellectual foundation of this work. "
        "Gratitude is also extended to researchers and practitioners in the Islamic "
        "digital humanities community for their feedback on system evaluation."
    ), size=12, space_after=12)

    divider(doc)
    heading(doc, "Data Availability", size=12, bold=True, space_before=10)
    para(doc, (
        f"The evaluation dataset (96 annotated queries), evaluation scripts, "
        f"and Kizana system source code are publicly available at: {GITHUB_URL}. "
        "The underlying corpus of 7,872 classical Arabic books is proprietary "
        "and not distributed; however, all evaluation data and code needed to "
        "reproduce the reported results are provided."
    ), size=12, space_after=0)

    path = out / "2_title_page_kbs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


# ============================================================
# 3. HIGHLIGHTS
# ============================================================
def create_highlights(out: Path):
    doc = Document()
    for section in doc.sections:
        section.left_margin  = Inches(1.25)
        section.right_margin = Inches(1.25)

    heading(doc, "Highlights", size=14, bold=True)
    para(doc, PAPER_TITLE, italic=True, size=11, space_after=10)
    divider(doc)
    para(doc, "Max 85 characters (incl. spaces) per bullet — KBS requirement:", italic=True, size=10, space_after=8)

    highlights = [
        "Kizana enables multilingual IR over 7,872 classical Arabic Islamic books",      # 74
        "Rule-based Islamic term knowledge base yields ΔMAP = +0.601 improvement",       # 71
        "MAP = 0.790 achieves 98.9% of Google Translate baseline (p = 0.148)",           # 69
        "Kizana outperforms Google Translate for Arabic queries (0.838 vs 0.831)",       # 71
        "Fully offline system: zero API dependency, deployable in pesantren settings",   # 76
    ]

    for i, h in enumerate(highlights, 1):
        char_count = len(h)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"• {h}")
        set_font(r, size=12)
        # Show character count as note
        r2 = p.add_run(f"  [{char_count}/85]")
        set_font(r2, size=9, italic=True, color=(120, 120, 120))

    path = out / "3_highlights_kbs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


# ============================================================
# 4. DECLARATION OF COMPETING INTERESTS  ("Author Agreement")
# ============================================================
def create_declaration_competing_interests(out: Path):
    """
    This is the 'Author Agreement' file KBS requires uploaded at submission.
    Generated to replicate the output of Elsevier's online declarations tool.
    Must be saved as .docx — author signature not required.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin   = Inches(1.5)
        section.right_margin  = Inches(1.5)

    # Header block — matches Elsevier declarations tool output style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DECLARATION OF COMPETING INTERESTS")
    set_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(JOURNAL_NAME)
    set_font(r, size=12, italic=True)
    p.paragraph_format.space_after = Pt(20)

    divider(doc)

    heading(doc, "Manuscript Title", size=11, bold=True, space_before=12, space_after=4)
    para(doc, PAPER_TITLE, size=12, italic=True, space_after=16)

    heading(doc, "Author(s)", size=11, bold=True, space_before=4, space_after=4)
    para(doc, AUTHOR_NAME, size=12, space_after=4)
    para(doc, AFFILIATION, size=11, italic=True, space_after=4)
    para(doc, AUTHOR_EMAIL, size=11, space_after=16)

    divider(doc)

    heading(doc, "Declaration", size=12, bold=True, space_before=12, space_after=8)

    para(doc, (
        "The author(s) declare(s) the following regarding competing interests "
        "in relation to the above-named manuscript:"
    ), size=12, space_after=10)

    # The formal declaration box
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    r = p.add_run("☑  I have nothing to declare.")
    set_font(r, size=13, bold=True)

    para(doc, "", space_after=6)

    para(doc, (
        "The author confirms there are no financial or personal relationships "
        "with other people or organizations that could inappropriately influence "
        "or bias this work, including but not limited to: employment, consultancies, "
        "stock ownership, honoraria, paid expert testimony, patent applications "
        "or registrations, grants, or other funding sources."
    ), size=11, space_after=16)

    divider(doc)

    heading(doc, "Funding", size=12, bold=True, space_before=10, space_after=6)
    para(doc, (
        "This research did not receive any specific grant from funding agencies "
        "in the public, commercial, or not-for-profit sectors."
    ), size=12, space_after=16)

    divider(doc)

    heading(doc, "Submission Declaration", size=12, bold=True, space_before=10, space_after=6)
    declarations = [
        "The work has not been published previously (except as an in-progress preprint).",
        "The article is not under consideration for publication elsewhere.",
        "The article's publication is approved by all named authors.",
        "If accepted, the article will not be published elsewhere in the same form without written consent of the copyright-holder.",
    ]
    for d in declarations:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(f"☑  {d}")
        set_font(r, size=11)

    para(doc, "", space_after=8)

    heading(doc, "Authorship Confirmation", size=12, bold=True, space_before=4, space_after=6)
    authorship_items = [
        "The author has made substantial contributions to the conception and design, acquisition of data, and analysis and interpretation of data.",
        "The author has drafted the article and revised it critically for important intellectual content.",
        "The author has given final approval of the version to be submitted.",
        "The author agrees to be accountable for all aspects of the work.",
    ]
    for item in authorship_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(f"☑  {item}")
        set_font(r, size=11)

    para(doc, "", space_after=20)
    divider(doc)

    heading(doc, "Signature", size=12, bold=True, space_before=10, space_after=8)
    para(doc, (
        "Note: Author signature is not required by Elsevier for this declaration. "
        "By submitting this document, the corresponding author confirms the above "
        "declarations on behalf of all listed authors."
    ), size=10, italic=True, space_after=12)

    para(doc, f"Name:  {AUTHOR_NAME}", bold=True, size=12, space_after=4)
    para(doc, f"Role:  Corresponding Author", size=12, space_after=4)
    para(doc, f"Date:  {DATE}", size=12, space_after=4)
    para(doc, f"Email: {AUTHOR_EMAIL}", size=12, space_after=0)

    path = out / "4_declaration_competing_interests_kbs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}  ← 'Author Agreement' file to upload")


# ============================================================
# 5. DECLARATION OF GENERATIVE AI USE
# ============================================================
def create_ai_declaration(out: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    heading(doc, "Declaration of Generative AI and AI-Assisted Technologies", size=13, bold=True)
    heading(doc, "in the Manuscript Preparation Process", size=13, bold=True, space_before=0, space_after=8)
    para(doc, PAPER_TITLE, italic=True, size=11, space_after=4)
    para(doc, JOURNAL_NAME, italic=True, size=11, space_after=12)
    divider(doc)

    para(doc, (
        "During the preparation of this work the author used GitHub Copilot "
        "(Microsoft/OpenAI, VS Code extension) in order to assist with code "
        "generation for the evaluation scripts, data processing pipelines, and "
        "LaTeX manuscript formatting. The AI tool was used to accelerate "
        "implementation of standard programming patterns; all research design, "
        "hypotheses, experimental protocols, statistical analyses, interpretation "
        "of results, and scientific conclusions were performed entirely by the author."
    ), size=12, space_after=10)

    para(doc, (
        "After using this tool, the author reviewed and edited all AI-assisted "
        "content as needed and takes full responsibility for the content of the "
        "published article. No AI tool was used to generate or alter figures, "
        "images, or artwork in this manuscript."
    ), size=12, space_after=10)

    para(doc, (
        "No AI tool was used for any part of scientific reasoning, literature "
        "synthesis, data analysis interpretation, or authoring of the core "
        "research findings. The Islamic jurisprudence terminology mappings, "
        "evaluation design, gold standard annotation, and all conclusions "
        "are entirely the author's original work."
    ), size=12, space_after=16)

    divider(doc)
    para(doc, f"Corresponding Author: {AUTHOR_NAME}", bold=True, size=12, space_after=4)
    para(doc, f"Date: {DATE}", size=12, space_after=0)

    path = out / "5_declaration_generative_ai_kbs.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")


# ============================================================
# MAIN
# ============================================================
def main():
    script_dir = Path(__file__).parent
    out = script_dir / "submission_kbs"
    out.mkdir(exist_ok=True)

    # Copy figures
    fig_src = script_dir / "submission" / "figures"
    fig_dst = out / "figures"
    if fig_dst.exists():
        shutil.rmtree(fig_dst)
    if fig_src.exists():
        shutil.copytree(fig_src, fig_dst)
        print(f"  ✓ figures/ ({len(list(fig_dst.iterdir()))} files)")

    # Copy compiled PDF manuscript (reuse from IP&M submission)
    pdf_src = script_dir / "submission" / "5_manuscript_anonymized.pdf"
    if pdf_src.exists():
        shutil.copy2(pdf_src, out / "6_manuscript_anonymized.pdf")
        print("  ✓ 6_manuscript_anonymized.pdf")

    # Copy LaTeX source
    tex_src = script_dir / "manuscript_anonymized.tex"
    if tex_src.exists():
        shutil.copy2(tex_src, out / "7_manuscript_source.tex")
        print("  ✓ 7_manuscript_source.tex")

    print("\nGenerating KBS submission documents...")

    create_cover_letter(out)
    create_title_page(out)
    create_highlights(out)
    create_declaration_competing_interests(out)
    create_ai_declaration(out)

    print(f"\n{'='*60}")
    print(f"KBS submission package ready: paper/submission_kbs/")
    print(f"{'='*60}")
    print("""
Upload order in KBS Editorial Manager:
  1_cover_letter_kbs.docx             → Cover Letter
  2_title_page_kbs.docx               → Title Page (with author details)
  3_highlights_kbs.docx               → Highlights
  4_declaration_competing_interests_kbs.docx → Competing Interests (Author Agreement)
  5_declaration_generative_ai_kbs.docx       → AI Declaration
  6_manuscript_anonymized.pdf         → Manuscript (no author details)
  7_manuscript_source.tex             → LaTeX Source File
  figures/Figure_1.pdf ... Figure_6.pdf → Figures (upload individually)

KBS-specific notes:
  • KBS uses NUMBER reference style [1], [2] — update manuscript if switching from IP&M
  • KBS prefers ≤20 double-spaced pages — your paper is ~23 pages, borderline OK
  • Data must be deposited: use https://github.com/dihannahdi/kizana
""")


if __name__ == "__main__":
    main()
