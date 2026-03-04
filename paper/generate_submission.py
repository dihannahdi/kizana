#!/usr/bin/env python3
"""
Kizana — IP&M Complete Submission Package Generator
=====================================================
Generates all required files for Elsevier IP&M journal submission:
  1. Cover Letter (Word)
  2. Title Page with Author Details (Word)
  3. Highlights (Word)
  4. Declaration of Competing Interests (Word)
  5. CRediT Author Statement (Word)
  6. Declaration of Generative AI (included in manuscript)

Usage:
    python generate_submission.py
"""

import os
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text, font_size=12, bold=False, italic=False, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, italic=italic)
    return p


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# ============================================================
# 1. COVER LETTER
# ============================================================
def create_cover_letter(output_dir: Path):
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    
    # Date
    add_paragraph(doc, "March 4, 2026", space_after=12)
    
    # Addressee
    add_paragraph(doc, "Dear Editor-in-Chief,", space_after=6)
    add_paragraph(doc, "Information Processing & Management", italic=True, space_after=12)
    
    # Subject
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Re: Submission of manuscript entitled ")
    set_font(run)
    run = p.add_run('"Kizana: Cross-Lingual Information Retrieval for Classical Islamic Jurisprudence Texts Using Domain-Specific Query Translation"')
    set_font(run, bold=True)
    
    # Body paragraphs
    paragraphs = [
        "We are pleased to submit the above-referenced manuscript for consideration as a Full Length Article in Information Processing & Management. This work presents Kizana, a novel cross-lingual information retrieval (CLIR) system that bridges modern Indonesian and English queries to a corpus of 7,872 classical Arabic Islamic texts—one of the world's largest pre-modern textual traditions.",
        
        "The key contributions of this manuscript are:",
    ]
    
    for text in paragraphs:
        add_paragraph(doc, text, space_after=6)
    
    # Bullet points
    contributions = [
        "A production-deployed CLIR system with a rule-based domain-specific query translation layer covering 200+ Islamic jurisprudence terms across four legal domains, enabling offline multilingual access to 3.4 million table-of-contents entries.",
        "A validated gold standard of 96 queries in four languages (Indonesian, English, Arabic, mixed) with inter-annotator agreement analysis (Cohen's κ_binary = 0.793, Krippendorff's α = 0.739).",
        "Empirical evidence that domain-specific rule-based translation achieves statistically equivalent performance to Google Translate + BM25 (MAP 0.790 vs. 0.799, p = 0.148), while operating entirely offline with zero API dependency.",
        "A systematic ablation study revealing that multi-variant expansion and Arabic morphological stemming—features commonly assumed to improve retrieval—actually degrade performance in this specialized domain."
    ]
    
    for c in contributions:
        p = doc.add_paragraph(c, style="List Bullet")
        for run in p.runs:
            set_font(run, size=12)
        p.paragraph_format.space_after = Pt(4)
    
    # More body
    more_paragraphs = [
        "This work is particularly relevant to IP&M's scope at the intersection of computing and information science, addressing a critical application in Islamic digital humanities with implications for information access in low-resource language pairs and specialized domains. The system serves a community of 250 million+ Indonesian Muslims who regularly consult classical Islamic texts for religious guidance.",
        
        "We believe this manuscript advances the state-of-the-art in domain-specific CLIR by demonstrating that curated terminology mappings can match neural machine translation for specialized retrieval tasks, with significant practical advantages in deployment constraints. The finding that less is sometimes more—that removing multi-variant expansion and stemming improves precision—challenges conventional assumptions in Arabic information retrieval.",
        
        "This manuscript has not been published previously and is not under consideration for publication elsewhere. All authors have approved the manuscript and agree with its submission to Information Processing & Management.",
        
        "We confirm that this work follows ethical guidelines and all data used in this research is publicly available classical Arabic text with no privacy concerns. The research did not involve human subjects beyond the planned user study protocol described in the paper.",
        
        "We look forward to reviewers' comments and appreciate your consideration of this manuscript.",
    ]
    
    for text in more_paragraphs:
        add_paragraph(doc, text, space_after=8)
    
    # Closing
    add_paragraph(doc, "Sincerely,", space_after=24)
    add_paragraph(doc, "Dihan Nahdi", bold=True, space_after=2)
    add_paragraph(doc, "Independent Researcher", space_after=2)
    add_paragraph(doc, "Kizana Search — bahtsulmasail.tech", italic=True, space_after=2)
    add_paragraph(doc, "Email: nahdi@bahtsulmasail.tech", space_after=2)
    
    filepath = output_dir / "1_cover_letter.docx"
    doc.save(str(filepath))
    print(f"  ✓ Cover letter: {filepath.name}")


# ============================================================
# 2. TITLE PAGE WITH AUTHOR DETAILS
# ============================================================
def create_title_page(output_dir: Path):
    doc = Document()
    
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Title
    add_paragraph(doc, "TITLE PAGE", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    
    # Article title
    add_paragraph(doc, "Article Title:", font_size=12, bold=True, space_after=6)
    add_paragraph(
        doc,
        "Kizana: Cross-Lingual Information Retrieval for Classical Islamic Jurisprudence Texts Using Domain-Specific Query Translation",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )
    
    # Authors
    add_paragraph(doc, "Author(s):", font_size=12, bold=True, space_after=6)
    add_paragraph(doc, "Dihan Nahdi", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # Affiliations
    add_paragraph(doc, "Affiliation(s):", font_size=12, bold=True, space_after=6)
    add_paragraph(doc, "Independent Researcher, Kizana Search Project", font_size=12,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_paragraph(doc, "https://bahtsulmasail.tech", font_size=12, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    # Corresponding author
    add_paragraph(doc, "Corresponding Author:", font_size=12, bold=True, space_after=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for label, value in [
        ("Name: ", "Dihan Nahdi"),
        ("\nEmail: ", "nahdi@bahtsulmasail.tech"),
        ("\nORCID: ", "[To be added]"),
    ]:
        run = p.add_run(label)
        set_font(run, bold=True)
        run = p.add_run(value)
        set_font(run)
    
    add_paragraph(doc, "", space_after=12)
    
    # Acknowledgements
    add_paragraph(doc, "Acknowledgements:", font_size=12, bold=True, space_after=6)
    add_paragraph(
        doc,
        "The author gratefully acknowledges the Islamic scholarly community (pesantren) "
        "for providing the domain expertise necessary for constructing the gold standard "
        "evaluation dataset and query translation dictionaries. The classical Arabic text "
        "corpus is sourced from publicly available digital libraries of Islamic heritage. "
        "The author thanks the anonymous reviewers for their constructive feedback.",
        space_after=18
    )
    
    # Declaration of competing interests
    add_paragraph(doc, "Declaration of Competing Interests:", font_size=12, bold=True, space_after=6)
    add_paragraph(
        doc,
        "The author declares that there are no known competing financial interests or "
        "personal relationships that could have appeared to influence the work reported "
        "in this paper.",
        space_after=18
    )
    
    # Funding
    add_paragraph(doc, "Funding:", font_size=12, bold=True, space_after=6)
    add_paragraph(
        doc,
        "This research did not receive any specific grant from funding agencies in the "
        "public, commercial, or not-for-profit sectors.",
        space_after=18
    )
    
    # CRediT Author Statement
    add_paragraph(doc, "CRediT Author Contribution Statement:", font_size=12, bold=True, space_after=6)
    
    roles = [
        ("Dihan Nahdi:", "Conceptualization, Methodology, Software, Validation, "
         "Formal Analysis, Investigation, Data Curation, Writing – Original Draft, "
         "Writing – Review & Editing, Visualization, Project Administration, Resources.")
    ]
    
    for author, contribution in roles:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(author + " ")
        set_font(run, bold=True)
        run = p.add_run(contribution)
        set_font(run)
    
    filepath = output_dir / "2_title_page.docx"
    doc.save(str(filepath))
    print(f"  ✓ Title page: {filepath.name}")


# ============================================================
# 3. HIGHLIGHTS
# ============================================================
def create_highlights(output_dir: Path):
    doc = Document()
    
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    add_paragraph(doc, "Highlights", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    highlights = [
        "Kizana enables cross-lingual search over 7,872 classical Arabic Islamic texts using Indonesian, English, or mixed-language queries through domain-specific rule-based translation.",
        "Domain-specific query translation achieves MAP = 0.790, reaching 98.9% of Google Translate + BM25 performance (MAP = 0.799) with no statistically significant difference (p = 0.148).",
        "Ablation study reveals query translation contributes ΔMAP = +0.601 while multi-variant expansion and Arabic stemming paradoxically degrade retrieval precision in this specialized domain.",
        "The system operates entirely offline with zero API dependency, making it deployable at resource-constrained Islamic educational institutions (pesantren) with limited internet access.",
        "Gold standard of 96 queries validated with substantial inter-annotator agreement (Cohen's κ = 0.793, Krippendorff's α = 0.739) across four languages and four Islamic legal domains."
    ]
    
    for h in highlights:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(h)
        set_font(run, size=12)
        p.paragraph_format.space_after = Pt(8)
    
    filepath = output_dir / "3_highlights.docx"
    doc.save(str(filepath))
    print(f"  ✓ Highlights: {filepath.name}")


# ============================================================
# 4. DECLARATION OF COMPETING INTERESTS
# ============================================================
def create_declaration_interests(output_dir: Path):
    doc = Document()
    
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    add_paragraph(doc, "Declaration of Competing Interests", font_size=14, bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    
    add_paragraph(
        doc,
        "The author declares that there are no known competing financial interests or "
        "personal relationships that could have appeared to influence the work reported "
        "in this paper.",
        space_after=12
    )
    
    add_paragraph(doc, "", space_after=24)
    add_paragraph(doc, "Dihan Nahdi", bold=True, space_after=2)
    add_paragraph(doc, "Date: March 4, 2026", space_after=2)
    
    filepath = output_dir / "4_declaration_competing_interests.docx"
    doc.save(str(filepath))
    print(f"  ✓ Declaration of interests: {filepath.name}")


# ============================================================
# ORGANIZE FIGURES
# ============================================================
def organize_figures(paper_dir: Path, output_dir: Path):
    figures_dir = output_dir / "figures"
    figures_dir.mkdir(exist_ok=True)
    
    fig_files = [
        "fig1_per_language_map",
        "fig2_ablation",
        "fig3_optimization",
        "fig4_per_domain",
        "fig5_system_comparison",
        "fig6_optimization_by_lang",
    ]
    
    count = 0
    for fig in fig_files:
        for ext in [".png", ".pdf"]:
            src = paper_dir / f"{fig}{ext}"
            if src.exists():
                # Rename to submission-friendly format
                dst = figures_dir / f"Figure_{count // 2 + 1}{ext}"
                shutil.copy2(str(src), str(dst))
                count += 1
    
    # Also copy with original names
    for fig in fig_files:
        for ext in [".png", ".pdf"]:
            src = paper_dir / f"{fig}{ext}"
            if src.exists():
                shutil.copy2(str(src), str(figures_dir / f"{fig}{ext}"))
    
    print(f"  ✓ Figures: {count} files copied to figures/")


# ============================================================
# MAIN
# ============================================================
def main():
    paper_dir = Path(__file__).parent
    output_dir = paper_dir / "submission"
    output_dir.mkdir(exist_ok=True)
    
    print("=" * 60)
    print("KIZANA — IP&M SUBMISSION PACKAGE GENERATOR")
    print("=" * 60)
    print(f"Output: {output_dir}\n")
    
    print("Generating submission files...")
    create_cover_letter(output_dir)
    create_title_page(output_dir)
    create_highlights(output_dir)
    create_declaration_interests(output_dir)
    organize_figures(paper_dir, output_dir)
    
    print("\n" + "=" * 60)
    print("SUBMISSION PACKAGE CONTENTS")
    print("=" * 60)
    print("""
Required files:
  1_cover_letter.docx              — Cover letter to editor
  2_title_page.docx                — Title page with author details,
                                      acknowledgements, CRediT, funding
  3_highlights.docx                — 5 key highlights (bullet points)
  4_declaration_competing_interests.docx — Declaration of interests
  manuscript_anonymized.pdf        — [Compile from LaTeX — see below]
  figures/                         — All figures (PNG + PDF, 300 DPI)

To compile the anonymized manuscript:
  Upload kizana_q1_paper.tex to VPS and compile with:
    xelatex kizana_q1_paper.tex
  Or use Overleaf (recommended for Arabic text support).
""")
    print("=" * 60)


if __name__ == "__main__":
    main()
